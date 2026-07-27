class IMAPConnectionError(Exception):
    """Raised when connecting to the IMAP server fails."""
    pass


class IMAPAuthenticationError(Exception):
    """Raised when logging in to the IMAP server fails."""
    pass


class IMAPMailboxError(Exception):
    """Raised when selecting or reading from the mailbox fails."""
    pass
import imaplib
import email
import re
import os
import io
import urllib.request
import urllib.parse
import zipfile
from email.header import decode_header

try:
    # pyrefly: ignore [missing-import]
    import pypdf
except ImportError:
    pypdf = None

try:
    # pyrefly: ignore [missing-import]
    import docx
except ImportError:
    docx = None

try:
    # pyrefly: ignore [missing-import]
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import xlrd
except ImportError:
    xlrd = None

def extract_text_from_attachment(content_bytes: bytes, filename: str) -> str:
    """
    Extracts text content from raw bytes of an attachment based on the file extension.

    Args:
        content_bytes: Raw bytes of the attachment
        filename: Name of the file (to determine extension)

    Returns:
        str: Extracted text content.
    """
    ext = filename.split('.')[-1].lower() if '.' in filename else ''

    if ext in ('txt', 'csv', 'json', 'xml', 'html', 'log', 'ini', 'cfg', 'md'):
        try:
            return content_bytes.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"[Error decoding text file: {e}]"

    elif ext == 'pdf':
        if pypdf is None:
            return "[Error: pypdf library is not installed for PDF extraction]"
        try:
            pdf_file = io.BytesIO(content_bytes)
            reader = pypdf.PdfReader(pdf_file)
            text = ""
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            return f"[Error extracting text from PDF: {e}]"

    elif ext in ('docx', 'doc'):
        if docx is None:
            return "[Error: python-docx library is not installed for Word document extraction]"
        try:
            docx_file = io.BytesIO(content_bytes)
            doc = docx.Document(docx_file)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text if cell.text is not None else "" for cell in row.cells]
                    text.append(" | ".join(row_text))
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from Word document: {e}]"

    elif ext == 'xlsx':
        if openpyxl is None:
            return "[Error: openpyxl library is not installed for Excel extraction]"
        try:
            excel_file = io.BytesIO(content_bytes)
            wb = openpyxl.load_workbook(excel_file, read_only=True, data_only=True)
            text = []
            for sheet in wb.worksheets:
                text.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = [str(cell) if cell is not None else "" for cell in row]
                        text.append(" | ".join(row_text))
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from Excel (.xlsx): {e}]"

    elif ext == 'xls':
        if xlrd is None:
            return "[Error: xlrd library is not installed for older Excel (.xls) extraction]"
        try:
            wb = xlrd.open_workbook(file_contents=content_bytes)
            text = []
            for sheet_idx in range(wb.nsheets):
                sheet = wb.sheet_by_index(sheet_idx)
                text.append(f"--- Sheet: {sheet.name} ---")
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    if any(cell != "" for cell in row):
                        row_text = [str(cell) if cell != "" else "" for cell in row]
                        text.append(" | ".join(row_text))
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from older Excel (.xls): {e}]"

    elif ext == 'zip':
        try:
            zip_file = io.BytesIO(content_bytes)
            text = []
            with zipfile.ZipFile(zip_file) as z:
                namelist = z.namelist()
                
                # Filter out directories, macOS metadata files, and hidden system files
                valid_files = []
                for name in namelist:
                    if name.endswith('/'):
                        continue
                    base_name = os.path.basename(name)
                    if name.startswith('__MACOSX/') or base_name.startswith('.') or base_name.startswith('._'):
                        continue
                    valid_files.append(name)
                
                text.append(f"--- Archive Contains {len(valid_files)} Readable/Extractable Files ---")
                for name in valid_files:
                    text.append(f"\n[File Inside Zip: {name}]")
                    try:
                        file_bytes = z.read(name)
                        file_text = extract_text_from_attachment(file_bytes, name)
                        indented_text = "\n".join("  " + line for line in file_text.splitlines())
                        text.append(indented_text)
                    except Exception as e:
                        text.append(f"  [Error extracting text: {e}]")
            return "\n".join(text).strip()
        except Exception as e:
            return f"[Error extracting text from Zip file: {e}]"


    else:
        return f"[Content extraction not supported for file type: .{ext}]"



def extract_drive_links(text_body, html_body):
    """
    Extracts Google Drive file links from text and HTML bodies.
    Returns:
        List[dict]: List of dicts with 'filename' and 'file_id'.
    """
    links = []
    # 1. From HTML Body
    if html_body:
        drive_pattern = r'href="https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)/[^"]*"[^>]*aria-label="([^"]+)"'
        matches = re.findall(drive_pattern, html_body)
        for file_id, filename in matches:
            filename = filename.strip()
            if filename and not any(item['file_id'] == file_id for item in links):
                links.append({'filename': filename, 'file_id': file_id})
        
        chip_pattern = r'href="https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)/[^*?"]*".*?class="[^"]*gmail_drive_chip[^"]*".*?<span dir="ltr"[^>]*>([^<]+)</span>'
        matches2 = re.findall(chip_pattern, html_body, re.DOTALL)
        for file_id, filename in matches2:
            filename = filename.strip()
            if filename and not any(item['file_id'] == file_id for item in links):
                links.append({'filename': filename, 'file_id': file_id})

    # 2. From Plain Text Body
    if text_body:
        lines = [line.strip() for line in text_body.splitlines() if line.strip()]
        for idx, line in enumerate(lines):
            match = re.search(r'drive\.google\.com/file/d/([a-zA-Z0-9_-]+)', line)
            if match:
                file_id = match.group(1)
                filename = "Google Drive File"
                if idx > 0 and lines[idx - 1] and not ("drive.google.com" in lines[idx - 1] or "http" in lines[idx - 1]):
                    filename = lines[idx - 1].strip("<>()[]\"' ")
                
                if filename and not any(item['file_id'] == file_id for item in links):
                    links.append({'filename': filename, 'file_id': file_id})
                    
    return links


def is_social_or_promotional_email(metadata_str: str, from_val: str, subject_val: str, msg_obj, is_gmail: bool) -> bool:
    """
    Determines if an email is a social notification or promotional bulk email.
    Checks Gmail system labels first (if Gmail), then falls back to header heuristics.
    """
    # 1. Check Gmail-specific category labels in metadata_str
    if is_gmail and metadata_str:
        if "\\Social" in metadata_str or "\\Promotions" in metadata_str:
            return True

    # 2. Check general headers
    from_header = (from_val or "").lower()
    subject_lower = (subject_val or "").lower()
    
    # Check List-Unsubscribe header
    has_unsubscribe = False
    is_bulk = False
    if msg_obj:
        has_unsubscribe = "list-unsubscribe" in [h.lower() for h in msg_obj.keys()]
        precedence = (msg_obj.get("Precedence") or "").lower()
        is_bulk = precedence in ("bulk", "list", "junk")
        
    # Social sender domains
    social_domains = ["facebookmail.com", "linkedin.com", "twitter.com", "instagram.com", "pinterest.com", "meetup.com", "social"]
    is_social_domain = any(domain in from_header for domain in social_domains)
    
    # Promotional keywords
    promo_keywords = ["promotion", "coupon", "discount", "deals", "newsletter", "special offer", "sale", "unsubscribe"]
    is_promo_keyword = any(kw in subject_lower or kw in from_header for kw in promo_keywords)
    
    if is_social_domain or (has_unsubscribe and (is_bulk or is_promo_keyword)):
        return True
        
    return False


def download_file_from_google_drive(file_id, dest_path=None):
    """
    Downloads a public file from Google Drive by its file ID.
    Handles the virus scan warning page automatically.
    """
    url = f"https://drive.google.com/uc?export=download&id={file_id}"
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    cookie_processor = urllib.request.HTTPCookieProcessor()
    opener = urllib.request.build_opener(cookie_processor)
    
    req = urllib.request.Request(url, headers=headers)
    with opener.open(req) as response:
        content = response.read()
        
        # Check if Google Drive returned a confirmation warning page
        html = content.decode('utf-8', errors='ignore') if content else ""
        if "confirm=" in html:
            match = re.search(r'confirm=([a-zA-Z0-9_-]+)', html)
            if match:
                confirm_token = match.group(1)
                confirm_url = f"https://drive.google.com/uc?export=download&confirm={confirm_token}&id={file_id}"
                req_confirm = urllib.request.Request(confirm_url, headers=headers)
                with opener.open(req_confirm) as confirm_response:
                    content = confirm_response.read()
                    html = content.decode('utf-8', errors='ignore') if content else ""

        # Verify if Google Drive redirected to a sign-in or account verification page
        if "accounts.google.com" in html and ("signin" in html or "ServiceLogin" in html or "base href" in html):
            raise PermissionError(
                f"File is private. Please open the link in your browser to download/view it:\n"
                f"  https://drive.google.com/file/d/{file_id}/view"
            )
                    
        if dest_path:
            with open(dest_path, "wb") as f:
                f.write(content)
            
        return content



def get_email_attachments_in_memory(host, username, password, msg_id, mailbox="INBOX"):
    """
    Connects to an IMAP server, fetches a specific email by msg_id,
    extracts all attachments (both MIME attachments and Google Drive links) directly into RAM,
    extracts text content, and returns a list of dictionaries with raw bytes and text content.
    """
    try:
        mail = imaplib.IMAP4_SSL(host)
    except Exception as e:
        raise IMAPConnectionError(f"Could not establish connection to host '{host}': {e}")

    try:
        try:
            mail.login(username, password)
        except imaplib.IMAP4.error as e:
            raise IMAPAuthenticationError(
                f"Failed to authenticate user '{username}' on IMAP server: {e}"
            )

        try:
            status, _ = mail.select(mailbox)
            if status != "OK":
                raise IMAPMailboxError(f"Unable to open mailbox: '{mailbox}'")
        except Exception as e:
            if not isinstance(e, IMAPMailboxError):
                raise IMAPMailboxError(f"Error selecting mailbox '{mailbox}': {e}")
            raise e

        # Fetch message by sequence ID / UID
        status, msg_data = mail.fetch(str(msg_id).encode(), "(RFC822)")
        if status != "OK" or not msg_data or not msg_data[0]:
            raise IMAPMailboxError(f"Failed to fetch message ID '{msg_id}' from mailbox.")

        msg = email.message_from_bytes(msg_data[0][1])
        attachments = []
        text_body = ""
        html_body = ""

        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        text_body += payload.decode('utf-8', errors='ignore')
                except:
                    pass
            elif content_type == "text/html":
                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_body += payload.decode('utf-8', errors='ignore')
                except:
                    pass

            filename = part.get_filename()
            content_disposition = part.get_content_disposition()
            
            # If it's a MIME attachment
            if filename or content_disposition == 'attachment':
                if not filename:
                    filename = "unnamed_attachment"
                
                # Decode filename
                decoded_filename_parts = decode_header(filename)
                decoded_filename = ""
                for decoded_text, encoding in decoded_filename_parts:
                    if isinstance(decoded_text, bytes):
                        try:
                            decoded_filename += decoded_text.decode(encoding or "utf-8", errors="ignore")
                        except Exception:
                            decoded_filename += decoded_text.decode("utf-8", errors="ignore")
                    else:
                        decoded_filename += decoded_text

                payload = part.get_payload(decode=True)
                if not payload:
                    continue

                # Extract text content directly from the in-memory payload bytes
                text_content = extract_text_from_attachment(payload, decoded_filename)

                attachments.append({
                    "filename": decoded_filename,
                    "content_bytes": payload,
                    "text_content": text_content
                })

        # Process Google Drive attachments as well
        drive_links = extract_drive_links(text_body, html_body)
        for item in drive_links:
            filename = item['filename']
            file_id = item['file_id']

            try:
                # Download Google Drive file directly to RAM
                payload = download_file_from_google_drive(file_id, dest_path=None)
                
                # Extract text content from the in-memory payload bytes
                text_content = extract_text_from_attachment(payload, filename)

                attachments.append({
                    "filename": filename,
                    "content_bytes": payload,
                    "text_content": text_content
                })
            except Exception as e:
                attachments.append({
                    "filename": filename,
                    "content_bytes": b"",
                    "text_content": f"[Error downloading Google Drive link: {e}]"
                })

        return attachments

    finally:
        try:
            mail.logout()
        except:
            pass


def download_email_attachments(host, username, password, msg_id, download_dir, mailbox="INBOX"):
    """
    Connects to an IMAP server, fetches a specific email by msg_id,
    extracts all attachments (both MIME attachments and Google Drive links),
    saves them to download_dir, and extracts text content.

    Args:
        host: IMAP server address
        username: Email address
        password: App password (plaintext decrypted)
        msg_id: IMAP message ID (as string)
        download_dir: Target directory path
        mailbox: Mailbox name

    Returns:
        List[dict]: Details of downloaded attachments including 'filename', 'saved_path', and 'text_content'.
    """
    try:
        mail = imaplib.IMAP4_SSL(host)
    except Exception as e:
        raise IMAPConnectionError(f"Could not establish connection to host '{host}': {e}")

    try:
        try:
            mail.login(username, password)
        except imaplib.IMAP4.error as e:
            raise IMAPAuthenticationError(
                f"Failed to authenticate user '{username}' on IMAP server: {e}"
            )

        try:
            status, _ = mail.select(mailbox)
            if status != "OK":
                raise IMAPMailboxError(f"Unable to open mailbox: '{mailbox}'")
        except Exception as e:
            if not isinstance(e, IMAPMailboxError):
                raise IMAPMailboxError(f"Error selecting mailbox '{mailbox}': {e}")
            raise e

        # Fetch message by sequence ID
        status, msg_data = mail.fetch(str(msg_id).encode(), "(RFC822)")
        if status != "OK" or not msg_data or not msg_data[0]:
            raise IMAPMailboxError(f"Failed to fetch message ID '{msg_id}' from mailbox.")

        msg = email.message_from_bytes(msg_data[0][1])
        downloaded = []

        # Create download directory if it does not exist
        if not os.path.exists(download_dir):
            os.makedirs(download_dir, exist_ok=True)

        text_body = ""
        html_body = ""

        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        text_body += payload.decode('utf-8', errors='ignore')
                except:
                    pass
            elif content_type == "text/html":
                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_body += payload.decode('utf-8', errors='ignore')
                except:
                    pass

            filename = part.get_filename()
            content_disposition = part.get_content_disposition()
            
            # If it's a MIME attachment
            if filename or content_disposition == 'attachment':
                if not filename:
                    filename = "unnamed_attachment"
                
                # Decode filename
                decoded_filename_parts = decode_header(filename)
                decoded_filename = ""
                for decoded_text, encoding in decoded_filename_parts:
                    if isinstance(decoded_text, bytes):
                        try:
                            decoded_filename += decoded_text.decode(encoding or "utf-8", errors="ignore")
                        except Exception:
                            decoded_filename += decoded_text.decode("utf-8", errors="ignore")
                    else:
                        decoded_filename += decoded_text

                payload = part.get_payload(decode=True)
                if not payload:
                    continue

                # Avoid overwriting files with conflict resolution (suffix)
                base_name, ext = os.path.splitext(decoded_filename)
                saved_path = os.path.join(download_dir, decoded_filename)
                counter = 1
                while os.path.exists(saved_path):
                    saved_path = os.path.join(download_dir, f"{base_name}_{counter}{ext}")
                    counter += 1

                with open(saved_path, "wb") as f:
                    f.write(payload)

                # Extract text content and store in a variable
                text_content = extract_text_from_attachment(payload, decoded_filename)

                downloaded.append({
                    "filename": os.path.basename(saved_path),
                    "saved_path": os.path.abspath(saved_path),
                    "text_content": text_content
                })

        # Process Google Drive attachments as well
        drive_links = extract_drive_links(text_body, html_body)
        for item in drive_links:
            filename = item['filename']
            file_id = item['file_id']

            # Resolve conflict in filenames
            base_name, ext = os.path.splitext(filename)
            saved_path = os.path.join(download_dir, filename)
            counter = 1
            while os.path.exists(saved_path):
                saved_path = os.path.join(download_dir, f"{base_name}_{counter}{ext}")
                counter += 1

            try:
                # Download Google Drive file
                payload = download_file_from_google_drive(file_id, saved_path)
                
                # Extract text content and store in a variable
                text_content = extract_text_from_attachment(payload, filename)

                downloaded.append({
                    "filename": os.path.basename(saved_path),
                    "saved_path": os.path.abspath(saved_path),
                    "text_content": text_content
                })
            except Exception as e:
                downloaded.append({
                    "filename": filename,
                    "saved_path": "[Failed to download Google Drive attachment]",
                    "text_content": f"[Error downloading Google Drive link: {e}]"
                })

        return downloaded


    finally:
        try:
            mail.logout()
        except:
            pass

